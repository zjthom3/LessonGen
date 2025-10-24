"""Services for exporting lessons to various formats."""
from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import datetime
from typing import Any

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from sqlalchemy.orm import Session

from app.models.lesson import Lesson, LessonVersion


@dataclass(slots=True)
class ExportContent:
    title: str
    subject: str
    grade_level: str
    objective: str | None
    duration_minutes: int | None
    standards: list[str]
    materials: list[str]
    flow: list[str]
    differentiation: list[str]
    assessments: list[str]
    accommodations: list[str]


class ExportService:
    """Handles conversion of lessons into PDF/DOCX/GDoc representations."""

    def __init__(self, session: Session) -> None:
        self.session = session

    def export(self, lesson: Lesson, version: LessonVersion, format_name: str) -> Any:
        content = self._build_content(lesson, version)
        format_name = format_name.lower()
        if format_name == "pdf":
            return self._export_pdf(content)
        if format_name == "docx":
            return self._export_docx(content)
        if format_name == "gdoc":
            return self._export_gdoc(content)
        raise ValueError("Unsupported export format")

    def _build_content(self, lesson: Lesson, version: LessonVersion) -> ExportContent:
        standards = [standard.code for standard in version.standards] if version.standards else []
        materials = [
            f"{item.get('label', 'Material')}: {item.get('value', '')}"
            for item in version.materials
        ]
        flow = [
            f"{step.get('phase', 'Activity')} ({step.get('minutes', '?')} min): {step.get('content_md', '')}"
            for step in version.flow
        ]
        differentiation = [
            f"{item.get('strategy', 'Support')}: {item.get('description', item)}"
            for item in version.differentiation
        ]
        assessments = [
            f"{item.get('type', 'Assessment')}: {item.get('description', item)}"
            for item in version.assessments
        ]
        accommodations = [
            f"{item.get('type', 'Accommodation')}: {item.get('description', item)}"
            for item in version.accommodations
        ]

        return ExportContent(
            title=lesson.title,
            subject=lesson.subject,
            grade_level=lesson.grade_level,
            objective=version.objective,
            duration_minutes=version.duration_minutes,
            standards=standards,
            materials=materials,
            flow=flow,
            differentiation=differentiation,
            assessments=assessments,
            accommodations=accommodations,
        )

    def _export_pdf(self, content: ExportContent) -> bytes:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.set_font("Helvetica", "B", 16)
        usable_width = pdf.w - pdf.l_margin - pdf.r_margin
        pdf.multi_cell(
            usable_width,
            10,
            content.title,
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(
            usable_width,
            8,
            f"Subject: {content.subject} (Grade {content.grade_level})",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        if content.duration_minutes:
            pdf.multi_cell(
                usable_width,
                8,
                f"Duration: {content.duration_minutes} minutes",
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
        pdf.ln(2)

        def add_section(title: str, lines: list[str] | str | None) -> None:
            if not lines:
                return
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(
                usable_width,
                8,
                title,
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
            pdf.set_font("Helvetica", size=11)
            if isinstance(lines, str):
                pdf.multi_cell(
                    usable_width,
                    6,
                    lines,
                    new_x=XPos.LMARGIN,
                    new_y=YPos.NEXT,
                )
            else:
                for line in lines:
                    pdf.multi_cell(
                        usable_width,
                        6,
                        f"- {line}",
                        new_x=XPos.LMARGIN,
                        new_y=YPos.NEXT,
                    )
            pdf.ln(1)

        add_section("Objective", content.objective or "")
        add_section("Standards", content.standards)
        add_section("Materials", content.materials)
        add_section("Lesson Flow", content.flow)
        add_section("Differentiation", content.differentiation)
        add_section("Assessments", content.assessments)
        add_section("Accommodations", content.accommodations)

        pdf_bytes = pdf.output()
        if isinstance(pdf_bytes, bytearray):
            return bytes(pdf_bytes)
        if isinstance(pdf_bytes, str):
            return pdf_bytes.encode("latin1", errors="replace")
        return pdf_bytes

    def _export_docx(self, content: ExportContent) -> bytes:
        document = Document()
        document.add_heading(content.title, level=1)
        document.add_paragraph(f"Subject: {content.subject}")
        document.add_paragraph(f"Grade Level: {content.grade_level}")
        if content.duration_minutes:
            document.add_paragraph(f"Duration: {content.duration_minutes} minutes")
        document.add_paragraph()

        def add_section(title: str, values: list[str] | str | None) -> None:
            if not values:
                return
            document.add_heading(title, level=2)
            if isinstance(values, str):
                document.add_paragraph(values)
            else:
                for entry in values:
                    document.add_paragraph(entry, style="List Bullet")

        add_section("Objective", content.objective or "")
        add_section("Standards", content.standards)
        add_section("Materials", content.materials)
        add_section("Lesson Flow", content.flow)
        add_section("Differentiation", content.differentiation)
        add_section("Assessments", content.assessments)
        add_section("Accommodations", content.accommodations)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _export_gdoc(self, content: ExportContent) -> dict[str, Any]:
        timestamp = datetime.utcnow().isoformat() + "Z"
        return {
            "title": content.title,
            "subject": content.subject,
            "grade_level": content.grade_level,
            "generated_at": timestamp,
            "sections": {
                "objective": content.objective,
                "standards": content.standards,
                "materials": content.materials,
                "flow": content.flow,
                "differentiation": content.differentiation,
                "assessments": content.assessments,
                "accommodations": content.accommodations,
            },
            "status": "ready",
            "message": "Stubbed Google Docs export. Upload manually to Google Docs interface.",
        }
